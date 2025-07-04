from crewai import LLM, Agent, Crew, Process, Task
from crewai_tools import FileReadTool, FileWriterTool
from dotenv import load_dotenv
from docx import Document
import fitz  # PyMuPDF
import io
import os

load_dotenv()


BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOADS_DIR = os.path.join(BASE_DIR, "test")
OUTPUTS_DIR = os.path.join(BASE_DIR, "test_output")

os.makedirs(UPLOADS_DIR, exist_ok=True) 
os.makedirs(OUTPUTS_DIR, exist_ok=True)

llm = LLM(
    model="gemini/gemini-2.5-flash-preview-04-17",
    api_key=os.getenv("GEMINI_API_KEY"),
    temperature=0.0,
    stream=True,
    reasoning_effort="high"
)


uploaded_srs = "C:\\Users\\l43ar\\Downloads\\mermaiddynamic\\srs.docx"
original_filename = "srs_output.docx"

def generate_sdd(uploaded_srs, original_filename=original_filename):

    if uploaded_srs:
        print("[INFO] Upload received")
        filename_base, ext = os.path.splitext(original_filename)
        ext = ext.lower()

        # Normalized .txt file path in the UPLOADS_DIR
        srs_path = os.path.join(UPLOADS_DIR, f"{filename_base}.txt")

        try:
            text_content = ""
            
            # Check if uploaded_srs is a file path (string) or file content (bytes)
            if isinstance(uploaded_srs, str):
                # It's a file path, read the file
                if ext == ".docx":
                    print("[INFO] Converting DOCX to text...")
                    doc = Document(uploaded_srs)
                    text_content = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                elif ext == ".pdf":
                    print("[INFO] Converting PDF to text...")
                    pdf = fitz.open(uploaded_srs)
                    text_content = "\n".join([page.get_text() for page in pdf])
                    pdf.close()
                elif ext == ".txt" or ext == ".md":
                    print("[INFO] Reading text file...")
                    with open(uploaded_srs, "r", encoding="utf-8") as f:
                        text_content = f.read()
                else:
                    raise ValueError(f"Unsupported file type: {ext}")
            else:
                # It's file content (bytes), process as before
                if ext == ".docx":
                    print("[INFO] Converting DOCX to text...")
                    doc = Document(io.BytesIO(uploaded_srs))
                    text_content = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                elif ext == ".pdf":
                    print("[INFO] Converting PDF to text...")
                    pdf = fitz.open(stream=io.BytesIO(uploaded_srs), filetype="pdf")
                    text_content = "\n".join([page.get_text() for page in pdf])
                    pdf.close()
                elif ext == ".txt" or ext == ".md":
                    text_content = uploaded_srs.decode('utf-8')
                else:
                    raise ValueError(f"Unsupported file type: {ext}")

            safe_text = ''.join([c if ord(c) < 128 or c in '\n\r\t' else '' for c in text_content])

            with open(srs_path, "w", encoding="utf-8") as f:
                f.write(safe_text)

            print(f"[INFO] Cleaned SRS content saved to {srs_path}")

        except Exception as e:
            print(f"[ERROR] Failed to process uploaded file: {e}")
            raise

    file_read_tool = FileReadTool(file_path=srs_path)
    sdd_output_filepath = os.path.join(OUTPUTS_DIR, f"{filename_base}_sdd.md")
    file_writer_tool = FileWriterTool()



    # Agents for SRS to SDD
    srs_extractor = Agent(
    role="SRS Extractor Agent",
    goal="Analyze the input SRS document and extract all the sections including key functional and non-functional details needed for system design.",
    backstory="An AI expert trained to read and interpret software requirement specifications and convert them into design-ready insights.",
    tools=[file_read_tool],
    verbose= True
    )

    architecture_designer = Agent(
    role="Architecture Designer",
    goal="Design the high-level and component architecture of the system using proper diagrams and structure.",
    backstory="A senior software architect responsible for breaking down the system into layers, components, and services based on the SRS.",
    verbose= True
    )

    data_modeler = Agent(
    role="Data Modeler",
    goal="Design the data flow diagrams, ER schema, and tabular data dictionary based on system features.",
    backstory="A data architect who builds visual and textual representations of the data models and storage design.",
    verbose=True

    )

    api_interface_designer = Agent(
    role="Interface/API Designer",
    goal="Define external/internal APIs and create tables with endpoint details. Also generate UI wireframe descriptions.",
    backstory="An experienced interface engineer who builds developer-friendly API documentation and UI layout strategies.",
    verbose=True

    )

    security_engineer = Agent(
    role="Security Engineer",
    goal="Identify security requirements and threats using STRIDE and recommend protection mechanisms.",
    backstory="An AI security analyst trained in authentication, authorization, and compliance standards.",
    verbose=True
    )

    quality_engineer = Agent(
    role="Quality Engineer",
    goal="Define non-functional requirements and quality attributes like scalability, performance, and reliability.",
    backstory="A senior quality architect specialized in ensuring systems are resilient, scalable, and efficient.",
    verbose=True
    )

    deployment_engineer = Agent(
    role="Deployment & Monitoring Engineer",
    goal="Design the CI/CD, deployment, monitoring, and infrastructure strategy using textual descriptions and diagrams.",
    backstory="An ops engineer responsible for reliable software deployment and observability systems.",
    verbose=True
    )

    integration_specialist = Agent(
    role="Integration Specialist",
    goal="Identify and illustrate third-party integration workflows and produce diagrams for integration flows.",
    backstory="An expert who ensures proper integration of external systems and APIs with the core application.",
    verbose=True
    )

    traceability_manager = Agent(
    role="Traceability Matrix Creator",
    goal="Map each functional requirement in the SRS to system components, design sections, and test cases in tabular format.",
    backstory="A specialist in documentation and traceability who ensures all requirements are linked across SDLC.",
    verbose=True
    )

    document_compiler = Agent(
    role="SDD Formatter & Compiler",
    goal="Assemble all outputs (text, tables, mermaid code, diagrams, wireframe layouts) into a single structured System Design Document in proper sequential order.",
    backstory="A master document generator AI skilled in formatting technical documents professionally.",
    tools=[file_writer_tool],
    verbose=True
    )



    # Define tasks
    extract_task = Task(
    description=f"Extract all major sections and subsections from the SRS file at {srs_path} such as: 1. Introduction, 2. System Overview, 3. Functional Requirements (with Word-style table of components, features, roles, inputs/outputs), 4. Non-Functional Requirements, 5. Data Management Strategy. Ensure there is no duplicate or overlapping content, and all extracted content is structured under correct headers.",
    expected_output="Structured text for Section 1 and 2 of SDD with accurate separation between functional and non-functional details. Include Functional Requirements table.",
    agent=srs_extractor,
    tools=[file_read_tool],
    llm=llm
    )

    architecture_task = Task(
    description=f"Generate system architecture based on the SRS file at {srs_path} including: 1. High-Level Architecture Diagram (Mermaid with styling, direction LR/TB, colors), 2. Component Diagram(Mermaid with styling, direction LR/TB, colors), 3. Technology Stack Table (Word-style), and 4. Design Patterns used. All Mermaid code should be styled and wrapped in proper markdown code blocks.",
    expected_output="Section 3 with architecture diagrams and word-style tabular tech stack using consistent section/subsection headers.",
    agent=architecture_designer,
    tools=[file_read_tool],
    llm=llm
    )

    data_design_task = Task(
    description=f"Create Section 5 based on the SRS file at {srs_path}: 1. Data Flow Diagrams (Level 0 and Level 1 using Mermaid with styling direction LR/TB, colors and labels), 2. ER Diagram in Mermaid with styled themes, and 3. Data Dictionary in Word-style tabular format with table name, column, type, constraints, and purpose.",
    expected_output="Section 5: Mermaid DFD, Mermaid ER diagram, and Word-style data dictionary table.",
    agent=data_modeler,
    tools=[file_read_tool],
    llm=llm
    )

    interface_task = Task(
    description="Define Section 6: 1. External/Internal Interface tables in Word-style format including API name, method, endpoint, input/output/auth. 2. Generate UI wireframe layout descriptions (not code) with panel sections like header, sidebar, content area, modal, and visual hierarchy for each screen. Clearly describe the layout styling.",
    expected_output="Section 6 with API tables and detailed layout-styled wireframe descriptions.",
    agent=api_interface_designer,
    llm=llm

    )

    security_task = Task(
    description="Generate Section 7 on security: 1. Threat modeling using STRIDE, 2. Mermaid threat diagrams with colors and layout annotations, 3. Describe authentication mechanisms (OAuth2, JWT, RBAC), data encryption, compliance needs (GDPR, HIPAA). Ensure threat diagram is stylized.",
    expected_output="Section 7: Textual description + Mermaid threat model with layout and visual enhancements.",
    agent=security_engineer,
    llm=llm
    )

    quality_task = Task(
    description="Define Section 8: 1. Non-functional requirements (like performance, uptime, latency), 2. Quality Attributes (scalability, maintainability, availability) in Word-style tabular format with column: Attribute, Description, Implementation Strategy.",
    expected_output="Section 8 with tabular NFRs and quality attributes.",
    agent=quality_engineer,
    llm=llm
    )

    workflow_state_task = Task(
    description="Design Section 9: 1. User journey flowcharts (in Mermaid with styling, layout direction, start-end points), 2. State Machine Diagrams using Mermaid. Use meaningful IDs, transitions, and wrap Mermaid code with styled code block markdown.",
    expected_output="Section 9 with activity and state machine diagrams (in Mermaid with styling).",
    agent=architecture_designer,
    llm=llm

    )

    deployment_task = Task(
    description="Generate Section 10 and 11: 1. Deployment architecture diagram (Mermaid with theme/style/infra nodes), 2. CI/CD flow explanation (tools, stages), 3. Monitoring strategy (tools like Grafana, ELK). All diagrams must be in Mermaid code block and stylized.",
    expected_output="Sections 10 and 11 with deployment diagrams and detailed text about monitoring setup.",
    agent=deployment_engineer,
    llm=llm
    )

    integration_task = Task(
    description="Generate Section 12: 1. External integrations with diagrams using Mermaid (styled sequence or flowcharts), 2. Description of each integration (API, webhook, trigger, fallback, retry logic). Mermaid code must include layout styling and comments.",
    expected_output="Section 12: Integration sequence diagrams + functional flow explanations.",
    agent=integration_specialist,
    llm=llm
    )

    traceability_task = Task(
    description="Generate Section 13: Traceability Matrix mapping each SRS functional requirement ID to the feature/component, linked design section, and related test case. Format this in a clean Word-style table with appropriate headers.",
    expected_output="Section 13: Traceability matrix in tabular format.",
    agent=traceability_manager,
    llm=llm
    )

    compile_task = Task(
    description="Compile all outputs into a fully structured System Design Document with ALL 14 sections: Introduction, System Overview, Functional Requirements, Non-Functional Requirements, Architecture, Components, Data Design, Interface, Security, Quality, Workflow, Deployment, Integration, Traceability, Conclusion, Appendix. Ensure Mermaid diagrams are wrapped in styled code blocks, all tables use Word-style formatting, wireframes have layout styling and descriptions, and no section is duplicated. Add a brief Conclusion and optional Appendix if applicable.",
    expected_output=f"Complete System Design Document (SDD) in structured sequence with styled diagrams, clean word-style tabular tables, mermaid code blocks with styling, layout-aware wireframes, and no section duplication. Also store the final output in the path {sdd_output_filepath}.",
    agent=document_compiler,
    llm=llm,
    output_file=sdd_output_filepath,
    tools=[file_writer_tool]
    )

        # Initialize and execute the Crew
    crew = Crew(
    agents=[
        srs_extractor,
        architecture_designer,
        data_modeler,
        api_interface_designer,
        security_engineer,
        quality_engineer,
        deployment_engineer,
        integration_specialist,
        traceability_manager,
        document_compiler
    ],
    tasks=[
        extract_task,
        architecture_task,
        data_design_task,
        interface_task,
        security_task,
        quality_task,
        workflow_state_task,
        deployment_task,
        integration_task,
        traceability_task,
        compile_task
    ],
    process=Process.sequential,
    verbose=True,
    llm=llm
    )
    
    crew.kickoff()

    if os.path.exists(sdd_output_filepath):
        with open(sdd_output_filepath, "r") as file:
            sdd_content = file.read()
        return sdd_content
    else:
        return "Error: File not found. Please Try Again!!"
    

if __name__ == "__main__":
    try:
        sdd_content = generate_sdd(uploaded_srs, original_filename)
        print("[INFO] SDD generation completed successfully.")
        print(sdd_content)  # Print or save the SDD content as needed
    except Exception as e:
        print(f"[ERROR] Failed to generate SDD: {e}")