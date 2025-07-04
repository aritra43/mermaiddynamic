import subprocess
import os
from docx import Document
from docx.shared import Inches

def chart_generator(mermaid_code, docx_path="Diagram_File_SDD_6.docx"):
    # Save Mermaid code to a .mmd file
    with open("diagram.mmd", "w") as f:
        f.write(mermaid_code)

    # Render Mermaid to PNG using Mermaid CLI
    try:
        subprocess.run([r"C:\Users\l43ar\AppData\Roaming\npm\mmdc.cmd", "-i", "diagram.mmd", "-o", "diagram.png"], check=True)
        print("Mermaid diagram generated successfully.")
    except subprocess.CalledProcessError:
        print("Error: Mermaid CLI failed to generate diagram.")
        return  # Exit if diagram generation fails

    # Check if Word Document exists, else create a new one
    if os.path.exists(docx_path):
        doc = Document(docx_path)
        print(f"Existing document found: {docx_path}. Appending diagram.")
    else:
        doc = Document()
        doc.add_heading('Document with Mermaid Diagrams', 0)
        print(f"No existing document found. Creating new document: {docx_path}")

    # Add diagram to the document
    doc.add_paragraph('The following diagram was generated using Mermaid:')
    doc.add_picture('diagram.png', width=Inches(5))

    # Save document
    doc.save(docx_path)
    print(f"Diagram appended to {docx_path}\n")
